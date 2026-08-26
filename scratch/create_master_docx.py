import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import os

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_styled(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level > 1 else 18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.bold = True
    if level == 1:
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(16, 44, 87) # Deep Navy
    elif level == 2:
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(30, 86, 160) # Royal Blue
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(50, 50, 50)
    return p

def add_callout(doc, title, text, bg_hex="F0F4F8", border_hex="1E56A0"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="36" w:space="0" w:color="{border_hex}"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"💡 {title}\n")
    r1.bold = True
    r1.font.name = 'Calibri'
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(16, 44, 87)
    
    r2 = p.add_run(text)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(40, 40, 40)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

print("Helper functions compiled successfully.")
