import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import os

doc = Document()

# Set standard margins
for section in doc.sections:
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_title(text, subtitle):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = RGBColor(16, 44, 87)
    
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(14)
    run2 = p2.add_run(subtitle)
    run2.font.name = 'Calibri'
    run2.font.size = Pt(12)
    run2.italic = True
    run2.font.color.rgb = RGBColor(100, 100, 100)

def add_h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(15)
    run.bold = True
    run.font.color.rgb = RGBColor(16, 44, 87)
    return p

def add_p(text, bold_prefix="", italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.bold = True
        r_pre.font.name = 'Calibri'
        r_pre.font.size = Pt(10.5)
        r_pre.font.color.rgb = RGBColor(30, 30, 30)
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(10.5)
    r.italic = italic
    r.font.color.rgb = RGBColor(40, 40, 40)
    return p

def add_bullet(text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.bold = True
        r_pre.font.name = 'Calibri'
        r_pre.font.size = Pt(10)
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(10)
    return p

def add_speech_box(time_range, title, spoken_lyrics, screen_action, source_file_lines, img_path=None):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.8)
    set_cell_background(cell, "F4F7FA")
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="12" w:space="0" w:color="1E56A0"/>
            <w:left w:val="single" w:sz="36" w:space="0" w:color="1E56A0"/>
            <w:bottom w:val="single" w:sz="12" w:space="0" w:color="1E56A0"/>
            <w:right w:val="single" w:sz="12" w:space="0" w:color="1E56A0"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    
    # Header
    r_hdr = p.add_run(f"⏱️ [{time_range}] — {title}\n")
    r_hdr.bold = True
    r_hdr.font.name = 'Calibri'
    r_hdr.font.size = Pt(12)
    r_hdr.font.color.rgb = RGBColor(16, 44, 87)
    
    # Screen Action
    r_act_hdr = p.add_run("🖥️ SCREEN ACTION / WHAT TO SHOW ON SCREEN:\n")
    r_act_hdr.bold = True
    r_act_hdr.font.name = 'Calibri'
    r_act_hdr.font.size = Pt(10)
    r_act_hdr.font.color.rgb = RGBColor(30, 86, 160)
    
    r_act = p.add_run(f"{screen_action}\n\n")
    r_act.font.name = 'Calibri'
    r_act.font.size = Pt(10)
    r_act.font.color.rgb = RGBColor(40, 40, 40)
    
    # Source File Lines
    r_file_hdr = p.add_run("📁 EXACT CODE FILE & LINE NUMBERS TO SHOW IN VS CODE:\n")
    r_file_hdr.bold = True
    r_file_hdr.font.name = 'Calibri'
    r_file_hdr.font.size = Pt(10)
    r_file_hdr.font.color.rgb = RGBColor(180, 50, 20)
    
    r_file = p.add_run(f"{source_file_lines}\n\n")
    r_file.font.name = 'Calibri'
    r_file.font.size = Pt(9.5)
    r_file.bold = True
    r_file.font.color.rgb = RGBColor(30, 30, 30)
    
    # Spoken Lyrics
    r_spk_hdr = p.add_run("🎙️ WORD-FOR-WORD SPOKEN SCRIPT (\"LYRICS\"):\n")
    r_spk_hdr.bold = True
    r_spk_hdr.font.name = 'Calibri'
    r_spk_hdr.font.size = Pt(10.5)
    r_spk_hdr.font.color.rgb = RGBColor(16, 44, 87)
    
    r_spk = p.add_run(f'"{spoken_lyrics}"\n')
    r_spk.font.name = 'Calibri'
    r_spk.font.size = Pt(10.5)
    r_spk.italic = True
    r_spk.font.color.rgb = RGBColor(20, 20, 20)
    
    # Image if available
    if img_path and os.path.exists(img_path):
        p_img = cell.add_paragraph()
        p_img.paragraph_format.space_before = Pt(6)
        p_img.paragraph_format.space_after = Pt(2)
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(img_path, width=Inches(6.0))
        
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_after = Pt(6)

# Title
add_title("PO Backlog Architect Agent — Word-for-Word Demo Script & Visual Timeline", 
          "Complete Minute-by-Minute Presentation Script ('Lyrics'), Visual Screen Actions, Code Line References, and Tab-by-Tab Demo Walkthrough for Digital T3 Submission")

add_p("This document is your exact step-by-step script for delivering your live presentation and video demo to Digital T3. It contains what to say ('lyrics'), what to display on screen at every moment, exact VS Code line numbers, and embedded UI screenshots.", bold_prefix="Presentation Setup Instructions: ")

add_bullet("Do NOT use PowerPoint slides. Share your desktop screen with Streamlit UI (http://localhost:8501) on the left half and VS Code on the right half.", "Visual Setup: ")
add_bullet("FastAPI Server running in Terminal 1 (uvicorn app.main:app --reload) and Streamlit running in Terminal 2 (streamlit run ui/dashboard.py).", "Pre-Demo Terminals: ")

img_dir = r"e:\Digital T3\po-backlog-architect-agent\docs\images"

# PART 1
add_h1("Part 1: 0:00 – 1:30 | Opening Pitch & Architecture Philosophy")
add_speech_box(
    time_range="0:00 – 1:30",
    title="The Opening Pitch & Governance Philosophy",
    spoken_lyrics=(
        "Hello everyone! Today I'm presenting the PO Backlog Architect Agent for FlowDesk. "
        "Notice in our Streamlit sidebar that we support two LLM providers via our abstract LLMProvider protocol: "
        "GroqProvider using open-weights model qwen/qwen3.6-27b for live AI inference, and MockProvider for fast, 0.05-second deterministic offline testing. "
        "When engineering teams try using LLMs to generate user stories and acceptance criteria, three major failure modes occur: "
        "First, hallucinated technical requirements—like an AI inventing a 50 MB file size limit when the specification has no exact number. "
        "Second, generic story proliferation—vague user stories like 'As a user, I want to manage data efficiently' that fail Definition of Ready. "
        "And third, uncontrolled external writes—where an LLM pushes unreviewed items directly into Jira or Azure DevOps as READY. "
        "Our core design philosophy is GOVERNANCE OVER GENERATION. "
        "The LLM proposes artifacts, but deterministic Python service code validates section-level citations, scores anti-generic specificity across 3 layers, "
        "evaluates Definition of Ready rules, and enforces a structural human approval gate with an immutable NOT_READY status floor."
    ),
    screen_action=(
        "1. Start on Streamlit UI homepage (http://localhost:8501). Point out sidebar dropdown: toggle between 'Groq (qwen/qwen3.6-27b)' and 'Mock (Deterministic)', and show active safeguards.\n"
        "2. Switch VS Code window to README.md lines 212-235 (Architecture ASCII Diagram showing LLMProvider interface)."
    ),
    source_file_lines="• README.md: L212 - L235 (Architecture Component Boundaries Diagram)\n• LLM Provider Base: app/llm/base.py: L1 - L30 (LLMProvider protocol)\n• Groq Live Provider: app/llm/groq_provider.py: L1 - L85 (qwen/qwen3.6-27b model)\n• Mock Offline Provider: app/llm/mock_provider.py: L1 - L75 (deterministic response mock)",
    img_path=os.path.join(img_dir, "tab1_context_search.png")
)

# PART 2
add_h1("Part 2: 1:30 – 3:00 | Context Indexing (O1) & Epic Decomposition (O2)")
add_speech_box(
    time_range="1:30 – 3:00",
    title="Context Indexing (O1) & Epic Decomposition (O2)",
    spoken_lyrics=(
        "First, let's look at Context Indexing in Tab 1. We parse the FlowDesk product brief into 15 addressable sections, indexed with SQLite FTS5. "
        "If I search 'file upload', the engine retrieves section PB-04.1. This addressable indexing is what enables claim-level citation validation.\n"
        "Now let's move to Tab 2: Epic Decomposition. If I select EP-001—a detailed document management epic—the agent generates structured stories with citations.\n"
        "However, look at EP-002—a thin, 6-word epic titled 'Automate Approval Overrides'. Instead of fabricating fake requirements, "
        "our decomposition agent detects that detail is missing and surfaces Open Questions to the Product Owner first, such as 'What role is authorized to perform overrides?'"
    ),
    screen_action=(
        "1. Click Streamlit Tab 1 (Context Search). Type 'file upload' in search bar, click 'Search Context'. Expand section card [PB-04.1].\n"
        "2. Click Streamlit Tab 2 (Epic Decomposer). Select 'EP-001', click 'Decompose Epic'. Show generated stories.\n"
        "3. Select 'EP-002' (Thin Epic), click 'Decompose Epic'. Point out the red warning banner and Open Questions surfaced."
    ),
    source_file_lines=(
        "• Context Indexer: app/services/context_service.py: L15 - L65 (index_markdown_brief)\n"
        "• Seed Epics: data/epics.json: L1 - L17 (EP-001 detailed vs EP-002 thin)\n"
        "• Decomposer Agent: app/agents/decomposition_agent.py: L20 - L80 (decompose_epic)"
    ),
    img_path=os.path.join(img_dir, "tab2_epic_decomposer.png")
)

# PART 3
add_h1("Part 3: 3:00 – 4:30 | Acceptance Criteria (O3), Citations (O6) & Anti-Generic Guard (O8)")
add_speech_box(
    time_range="3:00 – 4:30",
    title="Acceptance Criteria, Citation Validation & Anti-Generic Guard",
    spoken_lyrics=(
        "In Tab 3, we generate Given/When/Then acceptance criteria. When I select story BL-001, the agent produces testable GWT scenarios "
        "and attaches verified citation PB-04.1. Our CitationService checks both existence and numeric term grounding.\n"
        "Below, we have our 3-Layer Anti-Generic Guard. If someone inputs a vague requirement like 'Manage my data efficiently', "
        "our guard evaluates 3 layers: Layer 1 exact phrase match ('manage data'), Layer 2 vague verb regex ('manage'), and Layer 3 specificity score.\n"
        "It flags the story as GENERIC (Score = 1) and automatically rewrites it into a domain-specific user story with a measurable outcome."
    ),
    screen_action=(
        "1. Click Streamlit Tab 3 (Criteria Generator). Select story 'BL-001', click 'Generate Criteria'. Show GWT scenarios & PB-04.1 citation.\n"
        "2. Scroll down to Anti-Generic Guard. Type 'Manage my data efficiently' in custom text box, click 'Evaluate Specificity'.\n"
        "3. Point out 3-layer diagnostic breakdown (GENERIC, Score 1) and the auto-rewritten story."
    ),
    source_file_lines=(
        "• Citation Service: app/services/citation_service.py: L30 - L95 (validate_citation_support & numeric term check)\n"
        "• Generic Config: config/generic_guard.json: L1 - L35 (forbidden_phrases & vague_verbs)\n"
        "• Generic Guard Service: app/services/generic_guard_service.py: L40 - L110 (evaluate_specificity & rewrite)"
    ),
    img_path=os.path.join(img_dir, "tab3_criteria_generator.png")
)

# PART 4
add_h1("Part 4: 4:30 – 6:00 | Readiness Gate (O4), Prioritization (O5) & Overlap Detector (O7)")
add_speech_box(
    time_range="4:30 – 6:00",
    title="Definition of Ready Gate, Formula Prioritization & Overlap Detector",
    spoken_lyrics=(
        "In Tab 4, we evaluate the Definition of Ready. Selecting BL-003—which lacks acceptance criteria—fails rule check 'has_acceptance_criteria' "
        "and returns status BLOCKED. Selecting BL-005 passes all 6 YAML rules and returns status READY.\n"
        "In Tab 5, our Prioritization Engine calculates scores using a transparent mathematical formula combining Business Value, Urgency, Risk Reduction, "
        "Dependency Penalty, and Readiness Factor. We don't ask the LLM for priority; we compute it deterministically.\n"
        "In Tab 6, Overlap Detector checks new candidate story BL-006 against existing backlog stories and identifies it as a SUBSET of story BL-001 with 92% confidence, recommending a merge."
    ),
    screen_action=(
        "1. Click Streamlit Tab 4 (Readiness Gate). Select 'BL-003', click 'Evaluate Readiness' -> Show BLOCKED. Select 'BL-005' -> Show READY.\n"
        "2. Click Streamlit Tab 5 (Prioritization). Click 'Run Prioritization Engine' -> Show sorted priority table & formula breakdown.\n"
        "3. Click Streamlit Tab 6 (Overlap Detector). Select 'BL-006', click 'Check Overlap' -> Show SUBSET match with BL-001."
    ),
    source_file_lines=(
        "• Readiness YAML: config/readiness.yaml: L1 - L32 (6 active rules)\n"
        "• Readiness Service: app/services/readiness_service.py: L25 - L75 (evaluate_readiness)\n"
        "• Prioritization Service: app/services/prioritization_service.py: L30 - L80 (calculate_priority_score formula)\n"
        "• Overlap Service: app/services/overlap_service.py: L20 - L70 (check_overlap)"
    ),
    img_path=os.path.join(img_dir, "tab4_readiness_gate.png")
)

# PART 5
add_h1("Part 5: 6:00 – 7:30 | ⭐ The Killer Demo: Human Approval Gate & Status Floor (O9)")
add_speech_box(
    time_range="6:00 – 7:30",
    title="The Killer Demo: Human Approval Gate & Status Floor (O9)",
    spoken_lyrics=(
        "Now for our core safety control in Tab 7: The Human Approval Gate and Status Floor.\n"
        "Notice draft DFT-001 in the queue with status PENDING. If an external system or user attempts to write this unapproved draft directly to MockTracker, "
        "watch what happens when I click 'Write to Tracker': The service layer rejects the call and returns HTTP 403 Forbidden: Draft must be APPROVED.\n"
        "Now, as the human Product Owner, I click 'Approve Draft'. The status changes to APPROVED.\n"
        "Now I click 'Write to Tracker' again. The call succeeds with HTTP 200 OK. But look at MockTracker: "
        "Even though the draft was approved, the status is forcibly locked at NOT_READY and tagged 'AI-drafted'. "
        "This guarantees that no AI item enters a sprint without final human PO review."
    ),
    screen_action=(
        "1. Click Streamlit Tab 7 (Approval Queue). Select draft 'DFT-001' (Status: PENDING).\n"
        "2. FIRST CLICK: Click 'Write to Tracker' -> Show red error banner 'HTTP 403 Forbidden'. Open VS Code app/services/approval_service.py line 72.\n"
        "3. SECOND CLICK: Click 'Approve Draft' -> Show status badge update to APPROVED.\n"
        "4. THIRD CLICK: Click 'Write to Tracker' -> Show HTTP 200 OK success banner and MockTracker item with status NOT_READY and tag AI-drafted."
    ),
    source_file_lines=(
        "• Approval Service Gate: app/services/approval_service.py: L65 - L85 (line 72: if draft.status != ApprovalStatus.APPROVED: raise ApprovalRequiredError)\n"
        "• Status Floor Override: app/services/approval_service.py: L90 - L115 (write_draft_to_tracker: forces status=NOT_READY, tags=['AI-drafted'])"
    ),
    img_path=os.path.join(img_dir, "tab7_approval_queue.png")
)

# PART 6
add_h1("Part 6: 7:30 – 9:00 | Empirical Evaluation Harness & Adversarial Proofs")
add_speech_box(
    time_range="7:30 – 9:00",
    title="Empirical Evaluation Harness & Adversarial Probes",
    spoken_lyrics=(
        "To prove our system is robust beyond UI demonstrations, we built an automated evaluation suite.\n"
        "In Terminal, I run 'python -m eval.run --provider mock'—achieving 10/10 PASS across all Golden Cases in 0.05 seconds.\n"
        "Next, running 'python eval/compare.py' benchmarks our deterministic Mock against live Groq LLM inference with qwen/qwen3.6-27b over 3 repeated runs. "
        "Groq achieves 10/10 PASS with 1.45-second average latency, zero retries, and zero schema failures.\n"
        "Finally, running 'python eval/adversarial_run.py' executes 7 adversarial probes—achieving 7/7 PASS. "
        "For example, ADV-01 injects a hallucinated '50 MB' file limit. CitationService detects that '50 MB' is not in PB-04.2 and rejects the claim."
    ),
    screen_action=(
        "1. Switch screen focus to Terminal window.\n"
        "2. Type: python -m eval.run --provider mock -> Show 10/10 PASS output table.\n"
        "3. Type: python eval/compare.py -> Show side-by-side Mock vs Groq benchmark report.\n"
        "4. Type: python eval/adversarial_run.py -> Show 7/7 PASS output (pointing out ADV-01 numeric grounding rejection)."
    ),
    source_file_lines=(
        "• Golden Harness: eval/run.py: L20 - L100 (run_eval & provider CLI)\n"
        "• Side-by-Side Reporter: eval/compare.py: L10 - L60 (compare_results)\n"
        "• Adversarial Runner: eval/adversarial_run.py: L15 - L90 (run_adversarial_suite)"
    ),
    img_path=os.path.join(img_dir, "tab5_prioritization.png")
)

# PART 7
add_h1("Part 7: 9:00 – 10:00 | Architectural Trade-offs & Q&A Defense Strategy")
add_speech_box(
    time_range="9:00 – 10:00",
    title="Architectural Trade-offs & Q&A Defense",
    spoken_lyrics=(
        "In conclusion, our project evolved from a feature prototype into a governed AI Product Owner system.\n"
        "We made three key architectural decisions documented in our ADR log:\n"
        "First, we chose SQLite FTS5 over Vector DBs because our product brief is structured into addressable sections (PB-01 to PB-15). FTS5 provides 100% deterministic lookup with zero vector infrastructure overhead.\n"
        "Second, we built MockProvider alongside GroqProvider to enable fast, reproducible CI regression testing.\n"
        "And third, we enforced approval gating and status floors in Python service objects rather than complex graph frameworks, keeping the codebase clean, maintainable, and robust.\n"
        "Thank you! I am now ready for any technical questions."
    ),
    screen_action=(
        "1. Switch VS Code window to docs/decision-log.md (ADR-002, ADR-004, ADR-006, ADR-008).\n"
        "2. End presentation with full confidence."
    ),
    source_file_lines="docs/decision-log.md: L1 - L150 (ADR-001 through ADR-008 Architecture Decision Records)",
    img_path=os.path.join(img_dir, "tab6_overlap_detector.png")
)

# Save Document
out_doc_path = r"e:\Digital T3\po-backlog-architect-agent\docs\PO_Backlog_Architect_Demo_Presentation_Script.docx"
root_doc_path = r"e:\Digital T3\po-backlog-architect-agent\PO_Backlog_Architect_Demo_Presentation_Script.docx"

doc.save(out_doc_path)
doc.save(root_doc_path)

print(f"Demo presentation docx successfully generated at: {out_doc_path} and {root_doc_path}")
