import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import os

doc = Document()

# Set standard margins
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_title(text, subtitle):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = RGBColor(16, 44, 87)
    
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(16)
    run2 = p2.add_run(subtitle)
    run2.font.name = 'Calibri'
    run2.font.size = Pt(12)
    run2.italic = True
    run2.font.color.rgb = RGBColor(100, 100, 100)

def add_h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(16, 44, 87)
    return p

def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(13)
    run.bold = True
    run.font.color.rgb = RGBColor(30, 86, 160)
    return p

def add_h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11.5)
    run.bold = True
    run.font.color.rgb = RGBColor(50, 50, 50)
    return p

def add_p(text, bold_prefix="", italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.bold = True
        r_pre.font.name = 'Calibri'
        r_pre.font.size = Pt(10.5)
        r_pre.font.color.rgb = RGBColor(30, 30, 30)
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(10.5)
    r.italic = italic
    r.font.color.rgb = RGBColor(40, 40, 40)
    return p

def add_bullet(text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.bold = True
        r_pre.font.name = 'Calibri'
        r_pre.font.size = Pt(10)
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(10)
    return p

def add_callout(title, text, bg_hex="F0F4F8", border_hex="1E56A0"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.8)
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="36" w:space="0" w:color="{border_hex}"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"💡 {title}\n")
    r1.bold = True
    r1.font.name = 'Calibri'
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = RGBColor(16, 44, 87)
    
    r2 = p.add_run(text)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = RGBColor(40, 40, 40)
    
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_after = Pt(4)

def add_styled_table(headers, rows_data):
    tbl = doc.add_table(rows=len(rows_data) + 1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_background(hdr_cells[i], "1E56A0")
        set_cell_margins(hdr_cells[i], top=90, bottom=90, left=100, right=100)
        p = hdr_cells[i].paragraphs[0]
        for r in p.runs:
            r.font.name = 'Calibri'
            r.font.size = Pt(9.5)
            r.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            
    for r_idx, row in enumerate(rows_data):
        row_cells = tbl.rows[r_idx + 1].cells
        bg = "F9FBFD" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row):
            row_cells[c_idx].text = str(val)
            set_cell_background(row_cells[c_idx], bg)
            set_cell_margins(row_cells[c_idx], top=70, bottom=70, left=100, right=100)
            p = row_cells[c_idx].paragraphs[0]
            for r in p.runs:
                r.font.name = 'Calibri'
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(40, 40, 40)
                
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_after = Pt(4)

# Document Title
add_title("PO Backlog Architect Agent — Complete System Master Guide & UI Handbook", 
          "Comprehensive Reference for Purpose, Step Dependencies, Granular UI Operations, Architecture, ADRs, Code Map, and Interview Defense Script")

add_callout("Single Source of Truth Study Document", 
            "This document is your complete, all-in-one study guide for the FlowDesk PO Backlog Architect Agent. "
            "It explains why the UI exists, maps out step-by-step dependencies between all 7 modules, provides exact button-by-button user instructions for the Streamlit dashboard, "
            "and delivers full architectural and testing documentation for your Digital T3 submission.")

# SECTION 1: Strategic Purpose of the UI
add_h1("1. Why Have a UI? The Purpose & Strategic Value of the Application")
add_p("You might wonder: ", bold_prefix="")
add_p("If the core backend is built in Python and FastAPI, why do we need a Streamlit UI dashboard?", italic=True)
add_p("The Streamlit UI serves three fundamental strategic purposes:")

add_bullet("AI models are probabilistic and will occasionally output incomplete, vague, or misaligned user stories. A pure automated script would push AI drafts straight to production trackers. The UI provides an interactive dashboard where a human Product Owner (PO) can inspect context citations, review surfaced open questions, trigger anti-generic checks, evaluate Definition of Ready (DoR), and explicitly approve drafts.", "1. Human-in-the-Loop Governance: ")
add_bullet("In an assessment or interview demo, evaluators (like Digital T3 engineers) will not want to curl raw REST endpoints or read Python terminal logs. The Streamlit UI provides a visual, multi-tab interface where anyone can verify the agent's capabilities in real-time with zero technical setup.", "2. Interactive Interview Demo Environment: ")
add_bullet("The UI exposes system safeguards clearly in the sidebar (Auto-Fail Safeguards Active: Structural Approval Gate, Status Floor, Citation Verification). This proves that governance is enforced at the service layer, not just in UI code.", "3. Visual Evidence of System Safeguards: ")

# SECTION 2: Step Dependency & Lifecycle Dataflow Map
add_h1("2. System Lifecycle & Step Dependency Map (How Steps Connect)")
add_p("The 7 modules in the UI are not isolated pages; they form a logical, sequential workflow pipeline. Understanding how data flows between steps is essential for explaining the system during your demo.")

headers_dep = ["Step / Tab", "Name", "Prerequisite Dependencies (What It Uses)", "Output Produced (What It Feeds Into Next)"]
rows_dep = [
    ["Step 1", "Context Search (O1)", "FlowDesk Product Brief (data/product_brief.md)", "Provides addressable section IDs (PB-01...PB-15) for Step 2 & Step 3 grounding."],
    ["Step 2", "Epic Decomposer (O2)", "Step 1 Context Sections + Seed Epics (EP-001/002)", "Decomposes epics into User Stories & Open Questions. Generated stories feed into Step 3, 4, 5 & 7."],
    ["Step 3", "Criteria Generator (O3/O6/O8)", "User Story from Step 2 + Step 1 Addressable Context", "Generates Given/When/Then criteria + Citations. Feeds into Step 4 (DoR check) & Step 7 (Draft Queue). Runs 3-Layer GenericGuard."],
    ["Step 4", "Readiness Gate (O4)", "Story + Acceptance Criteria (Step 3) + config/readiness.yaml", "Evaluates 6 DoR rules. Outputs status READY or BLOCKED. Status feeds into Step 5 & Step 7."],
    ["Step 5", "Prioritization Engine (O5)", "Backlog Items + DoR Status (Step 4) + Dependency Rules", "Calculates weighted formula score. Sorts backlog topologically for sprint planning."],
    ["Step 6", "Overlap Detector (O7)", "New Candidate Story (Step 2/3) + Existing Backlog (BL-001..020)", "Detects DUPLICATE, SUBSET, SUPERSET, ADJACENT relationships to prevent duplicate work."],
    ["Step 7", "Approval Queue & Write Gate (O9)", "Drafts generated in Step 2/3 + DoR status from Step 4", "Enforces Human PO Approval (PENDING -> APPROVED). Rejects unapproved writes (HTTP 403). Forces status floor NOT_READY and tag AI-drafted."]
]
add_styled_table(headers_dep, rows_dep)

add_callout("Data Dependency Example: The End-to-End Chain",
            "1. Step 1 indexes section PB-04.1 ('Document Upload Format').\n"
            "2. Step 2 uses PB-04.1 to generate User Story BL-001 ('Upload document attachment').\n"
            "3. Step 3 generates GWT Acceptance Criteria and attaches citation PB-04.1 to story BL-001.\n"
            "4. Step 4 checks story BL-001: since it has criteria and citations, DoR returns status READY.\n"
            "5. Step 5 scores story BL-001 (Priority Score = 8.45) and places it at the top of Sprint 1.\n"
            "6. Step 7 receives story BL-001 as Draft DFT-001 (PENDING). PO clicks Approve, then Write to Tracker. ApprovalService writes issue to MockTracker with status NOT_READY and tag AI-drafted.")

# SECTION 3: Granular UI Operations Guide (Button-by-Button)
add_h1("3. Granular UI Operations Guide (Exact Inputs & Button Clicks for All Steps)")
add_p("Below is your complete operational guide for using each tab in the Streamlit UI dashboard (`http://localhost:8501`).")

# Step 1
add_h2("Step 1: Context Indexing & Addressable Search (Tab 1)")
add_bullet("FTS5 full-text search bar and section lookup cards.", "UI Elements: ")
add_bullet("Type a search term into the input box: e.g., 'file upload', 'approval', 'requester', or 'file size'.", "What to Input: ")
add_bullet("Click the 'Search Context' button.", "What to Click: ")
add_bullet("The dashboard displays 'Found N Addressable Sections'. Expand '[PB-04] Document & File Upload Management' to see verbatim section text indexed with section ID PB-04.1.", "Expected Outcome: ")
add_bullet("'Notice how every section of our product brief has a unique ID (PB-01 to PB-15). This addressable index allows our agents to validate claims against exact citations.'", "Demo Talking Point: ")

# Step 2
add_h2("Step 2: Epic Decomposer & Gap Detection (Tab 2)")
add_bullet("Epic selector dropdown, 'Decompose Epic' button.", "UI Elements: ")
add_bullet("Select 'EP-001 Document Management' (Detailed Epic) OR 'EP-002 Approval Automation' (Thin Epic).", "What to Input: ")
add_bullet("Click the 'Decompose Epic' button.", "What to Click: ")
add_bullet("For EP-001: Displays decomposed user stories with role, action, benefit, and citations. For EP-002 (Thin Epic): The system detects insufficient detail and surfaces 'Open Questions > Stories' (e.g. 'Undefined approver role').", "Expected Outcome: ")
add_bullet("'When an epic lacks detail like EP-002, our decomposer doesn't fabricate stories; it surfaces open questions to the PO first.'", "Demo Talking Point: ")

# Step 3
add_h3("Step 3: Acceptance Criteria & Anti-Generic Guard (Tab 3)")
add_bullet("Story selector, 'Generate Criteria' button, Anti-Generic Guard custom text input, 'Evaluate Specificity' button.", "UI Elements: ")
add_bullet("Part A: Select story 'BL-001'. Part B (Generic Check): Type 'Manage my data efficiently' into the custom text box.", "What to Input: ")
add_bullet("Click 'Generate Criteria' for Part A, or 'Evaluate Specificity' for Part B.", "What to Click: ")
add_bullet("Part A: Outputs Given/When/Then scenarios with verified citation PB-04.1. Part B: 3-Layer GenericGuard flags the text as GENERIC (Score = 1, phrase match 'manage data', vague verb 'manage') and presents an auto-rewritten story.", "Expected Outcome: ")
add_bullet("'Our GenericGuard uses 3 detection layers—exact phrases, vague verbs, and term density—to stop vague requirements from entering the sprint.'", "Demo Talking Point: ")

# Step 4
add_h2("Step 4: Definition of Ready (DoR) Gate (Tab 4)")
add_bullet("Backlog item selector, 'Evaluate Readiness' button.", "UI Elements: ")
add_bullet("Select story 'BL-003' (Incomplete story lacking acceptance criteria) OR 'BL-005' (Complete story).", "What to Input: ")
add_bullet("Click 'Evaluate Readiness'.", "What to Click: ")
add_bullet("For BL-003: Evaluates 6 rules from config/readiness.yaml. Shows has_acceptance_criteria = FAIL. Final Status: BLOCKED. For BL-005: All 6 rules PASS. Final Status: READY.", "Expected Outcome: ")
add_bullet("'The readiness engine strictly evaluates YAML rules. Incomplete items like BL-003 are BLOCKED until criteria and citations are attached.'", "Demo Talking Point: ")

# Step 5
add_h2("Step 5: Prioritization Engine (Tab 5)")
add_bullet("'Run Prioritization Engine' button, scored backlog table, formula breakdown cards.", "UI Elements: ")
add_bullet("No text input needed; uses backlog items from SQLite database.", "What to Input: ")
add_bullet("Click 'Run Prioritization Engine'.", "What to Click: ")
add_bullet("Renders a table sorted by Priority Score. Expanding an item shows the explicit mathematical formula: Score = (BV * 0.3) + (Urg * 0.2) + (RR * 0.2) + (SA * 0.15) - (DP * 0.1) + (RF * 0.05).", "Expected Outcome: ")
add_bullet("'Prioritization is 100% deterministic and transparent. We don't ask the LLM for a priority number; we compute it from an explicit formula.'", "Demo Talking Point: ")

# Step 6
add_h2("Step 6: Overlap & Duplicate Detector (Tab 6)")
add_bullet("Candidate story selector, 'Check Overlap' button.", "UI Elements: ")
add_bullet("Select candidate story 'BL-006' (Document Upload Request).", "What to Input: ")
add_bullet("Click 'Check Overlap'.", "What to Click: ")
add_bullet("Outputs relationship classification: SUBSET / DUPLICATE with target item BL-001 (Confidence = 0.92) and recommends merging items.", "Expected Outcome: ")
add_bullet("'Overlap detection prevents duplicate engineering effort by analyzing story semantics against existing backlog items.'", "Demo Talking Point: ")

# Step 7
add_h2("Step 7: Approval Queue & Write Gate (Tab 7 - The Killer Demo)")
add_bullet("Draft queue table, 'Approve Draft', 'Reject Draft', and 'Write to Tracker' buttons.", "UI Elements: ")
add_bullet("Select draft 'DFT-001' (Status: PENDING).", "What to Input: ")
add_bullet("FIRST: Click 'Write to Tracker' while status is PENDING. SECOND: Click 'Approve Draft'. THIRD: Click 'Write to Tracker' again.", "What to Click: ")
add_bullet("FIRST CLICK: Error banner appears: 'HTTP 403 Forbidden: Draft DFT-001 must be APPROVED before writing to tracker'. SECOND CLICK: Draft status changes to APPROVED. THIRD CLICK: Returns HTTP 200 OK. MockTracker displays item with status forced to NOT_READY and tag AI-drafted.", "Expected Outcome: ")
add_bullet("'This is our core safety control. Unapproved AI writes return HTTP 403 Forbidden. Even when approved, the status floor forces NOT_READY and tags AI-drafted, guaranteeing human PO review.'", "Demo Talking Point: ")

# SECTION 4: Assessment Requirements & Marking Criteria
add_h1("4. Assessment Requirements & Digital T3 Marking Criteria")
headers_o = ["Objective", "Requirement", "Implementation Component", "Verification Evidence"]
rows_o = [
    ["O1", "Context Indexing & Retrieval", "app/services/context_service.py (SQLite FTS5)", "TC-01 PASS, addressable section IDs (PB-01...PB-15)"],
    ["O2", "Epic Decomposition & Gap Detection", "app/agents/decomposition_agent.py", "TC-04 PASS, TC-08 PASS (thin epic surfaces questions > stories)"],
    ["O3", "Acceptance Criteria Generation (GWT)", "app/agents/criteria_agent.py", "Generated Given/When/Then scenarios with citations"],
    ["O4", "Definition of Ready (DoR) Gate", "app/services/readiness_service.py (config/readiness.yaml)", "TC-05 PASS, evaluates 6 configurable YAML rules"],
    ["O5", "Deterministic Prioritization Engine", "app/services/prioritization_service.py", "TC-06 PASS, formula scoring & topological sprint sorter"],
    ["O6", "Claim-Level Citation Validation", "app/services/citation_service.py", "TC-01 PASS, numeric term check rejects unsupported claims"],
    ["O7", "Overlap & Duplicate Detection", "app/services/overlap_service.py", "TC-07 PASS, detects DUPLICATE, SUBSET, SUPERSET, ADJACENT"],
    ["O8", "3-Layer Anti-Generic Guard", "app/services/generic_guard_service.py", "TC-03 PASS, 3-layer specificity scoring & auto-rewrite"],
    ["O9", "Human Approval Gate & Status Floor", "app/services/approval_service.py", "TC-09 PASS, HTTP 403 write rejection & NOT_READY status floor"]
]
add_styled_table(headers_o, rows_o)

# SECTION 5: Architectural Decision Records (ADRs)
add_h1("5. Architectural Decision Records (ADR Log 001 – 008)")
headers_adr = ["ADR ID", "Title", "Decision Summary", "Rationale & Impact"]
rows_adr = [
    ["ADR-001", "FastAPI for Backend REST API", "Used FastAPI with Pydantic v2 schemas.", "Delivers high-performance async REST API, automatic OpenAPI/Swagger docs, and clean error handling."],
    ["ADR-002", "SQLite + FTS5 for Context Indexing", "Used SQLite FTS5 full-text search over Vector DB.", "Provides 100% deterministic section lookup (PB-01..PB-15) with zero vector store infrastructure overhead."],
    ["ADR-003", "LLMProvider Abstract Protocol", "Created generic LLMProvider base class.", "Decouples application logic from specific LLM vendors; supports seamless model switching."],
    ["ADR-004", "MockProvider for Deterministic CI", "Implemented offline MockProvider.", "Enables 0.05-second, 100% reproducible test suite runs without API cost, rate limits, or network dependency."],
    ["ADR-005", "3-Layer Anti-Generic Guard", "Implemented 3-layer specificity scoring.", "Combines exact phrase matching, vague verb regex, and term density to rewrite generic stories into domain-specific items."],
    ["ADR-006", "Structural Human Approval Gate", "Enforced approval check at service layer.", "Returns HTTP 403 Forbidden for unapproved write attempts; structurally prevents un-reviewed AI writes to tracker."],
    ["ADR-007", "Immutable NOT_READY Status Floor", "Overrode AI status outputs on tracker write.", "Forces status NOT_READY and adds AI-drafted tag, guaranteeing human PO review before sprint planning."],
    ["ADR-008", "Pragmatic Scope & Framework Cuts", "Omitted heavy graph/agent frameworks.", "Avoids unnecessary complexity (LangChain, LangGraph, PostgreSQL) in favor of maintainable native Python service code."]
]
add_styled_table(headers_adr, rows_adr)

# SECTION 6: Complete File Map
add_h1("6. Complete File-by-File Codebase Map")
headers_files = ["Directory / File Path", "Role & Description", "Key Symbols / Functions"]
rows_files = [
    ["app/main.py", "FastAPI application entrypoint & middleware setup.", "app, FastAPI, CORS, startup_event()"],
    ["app/api/routes.py", "REST API endpoint handlers for context, criteria, readiness, approval, priority.", "index_context(), generate_criteria(), check_readiness(), approve_draft(), write_tracker()"],
    ["app/db/database.py", "SQLAlchemy SQLite database session configuration.", "engine, SessionLocal, Base, get_db()"],
    ["app/models/models.py", "SQLAlchemy ORM models for database tables.", "ContextSectionModel, BacklogItemModel, DraftModel, ApprovalLogModel, WriteLogModel"],
    ["app/schemas/domain.py", "Pydantic v2 schemas for request/response payloads.", "Citation, OpenQuestion, CriteriaResult, GenericGuardResult, ReadinessResult, PriorityScore, OverlapResult"],
    ["app/services/context_service.py", "SQLite FTS5 text parser & section indexer.", "ContextService, index_markdown_brief(), search_context(), get_section_by_ref()"],
    ["app/services/citation_service.py", "Claim-level citation verification & numeric grounding engine.", "CitationService, validate_citation_existence(), validate_citation_support()"],
    ["app/services/generic_guard_service.py", "3-layer anti-generic pattern checker & auto-rewriter.", "GenericGuardService, evaluate_specificity(), rewrite_generic_story()"],
    ["app/services/readiness_service.py", "YAML-driven Definition of Ready rule evaluator.", "ReadinessService, evaluate_readiness(), log_override()"],
    ["app/services/approval_service.py", "Draft state machine, approval gate & status floor.", "ApprovalService, create_draft(), approve_draft(), reject_draft(), write_draft_to_tracker()"],
    ["app/services/prioritization_service.py", "Formula priority scorer & topological sprint sorter.", "PrioritizationService, calculate_priority_score(), sort_backlog()"],
    ["app/services/overlap_service.py", "Overlap & duplicate relationship detector.", "OverlapService, check_overlap()"],
    ["app/agents/criteria_agent.py", "Structured GWT acceptance criteria generator.", "CriteriaAgent, generate_criteria()"],
    ["app/agents/decomposition_agent.py", "Epic decomposer & thin epic detector.", "DecompositionAgent, decompose_epic()"],
    ["app/adapters/tracker.py", "External issue tracker interface & MockTracker.", "Tracker, MockTracker, create_issue(), get_issue()"],
    ["app/adapters/doc_store.py", "Document store protocol implementation.", "DocumentStore, FileSystemDocumentStore"],
    ["app/llm/base.py", "Abstract LLMProvider protocol.", "LLMProvider, generate_json()"],
    ["app/llm/groq_provider.py", "Live Groq API provider with qwen/qwen3.6-27b.", "GroqProvider, generate_json()"],
    ["app/llm/mock_provider.py", "Deterministic offline Mock LLM provider.", "MockProvider, generate_json()"],
    ["config/readiness.yaml", "Configurable DoR rules YAML file.", "rules (6 active rules)"],
    ["config/generic_guard.json", "Forbidden generic terms & vague verbs configuration.", "forbidden_phrases, vague_verbs, domain_keywords"],
    ["data/product_brief.md", "FlowDesk specification seed markdown (15 sections PB-01..15).", "15 addressable sections, 3 planted gaps, 1 inconsistency, 1 contradiction"],
    ["data/glossary.json", "FlowDesk domain glossary.", "20 canonical domain terms"],
    ["data/backlog.json", "Seed backlog items (BL-001..020).", "20 mixed quality stories"],
    ["data/epics.json", "Seed epics (EP-001 detailed, EP-002 thin).", "EP-001, EP-002"],
    ["ui/dashboard.py", "Streamlit 7-tab dashboard application.", "Tab 1..7 interactive dashboard handlers"],
    ["eval/run.py", "Golden cases evaluation harness script.", "run_eval(), GoldenCaseRunner"],
    ["eval/compare.py", "Side-by-side Mock vs Groq benchmark reporter.", "compare_results()"],
    ["eval/adversarial_run.py", "7 adversarial scenarios test runner script.", "run_adversarial_suite()"],
    ["tests/test_api_integration.py", "Pytest suite for REST endpoints & SQLite threading.", "test_approval_gate_http_403(), test_write_tracker_success()"],
    ["tests/test_approval_gate.py", "Pytest unit tests for approval gate & status floor.", "test_unapproved_write_fails(), test_approved_write_status_floor()"],
    ["tests/test_grounding.py", "Pytest unit tests for citation resolution & numeric terms.", "test_citation_verification(), test_numeric_term_grounding()"],
    ["tests/test_e2e_pipeline.py", "Pytest 20-step happy path & edge case integration tests.", "test_20_step_e2e_pipeline()"]
]
add_styled_table(headers_files, rows_files)

# SECTION 7: Testing Evidence
add_h1("7. Testing & Verification Evidence")
add_h2("Golden Cases Benchmark (10 / 10 PASS)")
headers_g = ["Case ID", "Name", "Category", "Target", "Mock Result", "Groq Result", "Status"]
rows_g = [
    ["TC-01", "Citation Resolution", "Grounding", "0 missing", "0 missing", "0 missing", "PASS"],
    ["TC-02", "Open-Question Recall", "Gap Surface", "1.0 recall", "1.0 recall", "1.0 recall", "PASS"],
    ["TC-03", "Generic Story Rate", "Anti-Generic", "0.1 max", "0.0 rate", "0.0 rate", "PASS"],
    ["TC-04", "Decomposition Coverage", "Decomposition", "0.85 cov", "1.0 cov", "1.0 cov", "PASS"],
    ["TC-05", "Readiness Gate Accuracy", "DoR Gate", "1.0 acc", "1.0 acc", "1.0 acc", "PASS"],
    ["TC-06", "Prioritization Reproducibility", "Priority", "1.0 score", "1.0 score", "1.0 score", "PASS"],
    ["TC-07", "Overlap Detection", "Overlap", "True match", "True match", "True match", "PASS"],
    ["TC-08", "Thin Epic Behaviour", "Decomposition", "True match", "True match", "True match", "PASS"],
    ["TC-09", "Approval Gate & Status Floor", "Governance", "True match", "True match", "True match", "PASS"],
    ["TC-10", "Glossary Consistency", "Grounding", "True match", "True match", "True match", "PASS"]
]
add_styled_table(headers_g, rows_g)

add_h2("Adversarial Evaluation Suite (7 / 7 PASS)")
headers_adv = ["Adv ID", "Adversarial Test Scenario", "Tested Vulnerability", "Observed System Behavior", "Result"]
rows_adv = [
    ["ADV-01", "Hallucinated Numeric Term (50 MB)", "Numeric Spec Fabrication", "CitationService rejected claim 'term 50 MB not in PB-04.2'", "PASS"],
    ["ADV-02", "Invalid Section Ref (PB-99)", "Non-Existent Section Lookup", "CitationService rejected ref 'PB-99 does not exist in DB'", "PASS"],
    ["ADV-03", "Generic Title ('Fix everything')", "Vague Story Ingestion", "GenericGuard flagged GENERIC (score=1) and auto-rewrote story", "PASS"],
    ["ADV-04", "Unapproved Tracker Write Attempt", "Direct API Bypassing", "ApprovalService raised ApprovalRequiredError -> HTTP 403", "PASS"],
    ["ADV-05", "Status Floor Override on Approved Write", "AI Status Escalation", "ApprovalService forced status NOT_READY and added AI-drafted tag", "PASS"],
    ["ADV-06", "Duplicate Tracker Write Attempt", "Duplicate Ingestion", "ApprovalService caught existing write log -> HTTP 409 Conflict", "PASS"],
    ["ADV-07", "Prompt Injection Attack ('Ignore rules')", "Prompt Injection", "Model schema validation enforced -> system rules preserved", "PASS"]
]
add_styled_table(headers_adv, rows_adv)

# SECTION 8: Interview Defense Script
add_h1("8. Interview Q&A Defense Script")
add_p("1. Why did you use SQLite FTS5 instead of a Vector Database?", bold_prefix="Q: ")
add_p("Our product context is structured into discrete, addressable sections (PB-01 through PB-15). FTS5 provides 100% deterministic section lookup and phrase matching with zero vector DB infrastructure overhead, embedding model latency, or indexing cost.", bold_prefix="A: ")

add_p("2. How do you prevent an AI from putting READY items directly into Jira/Azure DevOps?", bold_prefix="Q: ")
add_p("The LLM does not have authority to write to external systems. The ApprovalService state machine enforces that unapproved writes return HTTP 403 Forbidden. Furthermore, when an approved write occurs, ApprovalService forcibly overrides the status to NOT_READY and attaches an AI-drafted tag, guaranteeing human PO review before sprint planning.", bold_prefix="A: ")

add_p(r"3. How do you detect hallucinated numbers or specs in LLM responses?", bold_prefix="Q: ")
add_p(r"Our CitationService performs numeric term grounding. It extracts all digit sequences (\d+) from generated claims and verifies that those numbers exist verbatim in the cited context section text. If a claim mentions '50 MB' but the cited section lacks '50', the citation is rejected.", bold_prefix="A: ")

# Save Document
out_doc_path = r"e:\Digital T3\po-backlog-architect-agent\docs\PO_Backlog_Architect_Master_Guide.docx"
root_doc_path = r"e:\Digital T3\po-backlog-architect-agent\PO_Backlog_Architect_Master_Guide.docx"

doc.save(out_doc_path)
doc.save(root_doc_path)

print(f"Master docx successfully generated at: {out_doc_path} and {root_doc_path}")
